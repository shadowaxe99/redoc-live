import PyPDF2
from docx import Document

DOCUMENT_FORMATS = ['pdf', 'docx', 'txt']

def parseDocument(file_path):
    file_extension = file_path.split('.')[-1]

    if file_extension not in DOCUMENT_FORMATS:
        raise ValueError(f"Unsupported file format. Supported formats are {DOCUMENT_FORMATS}")

    if file_extension == 'pdf':
        return parsePDF(file_path)
    elif file_extension == 'docx':
        return parseDOCX(file_path)
    elif file_extension == 'txt':
        return parseTXT(file_path)

def parsePDF(file_path):
    pdf_file_obj = open(file_path, 'rb')
    pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
    num_pages = pdf_reader.numPages
    document_content = ''

    for page in range(num_pages):
        page_obj = pdf_reader.getPage(page)
        document_content += page_obj.extractText()

    pdf_file_obj.close()
    return document_content

def parseDOCX(file_path):
    doc = Document(file_path)
    document_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return document_content

def parseTXT(file_path):
    with open(file_path, 'r') as file:
        document_content = file.read()
    return document_content
```
import logging
from datetime import datetime

logging.basicConfig(filename='app.log', level=logging.INFO, format='%(asctime)s - %(message)s', datefmt='%d-%b-%y %H:%M:%S')

def log_event(event):
    logging.info(event)

def log_exception(exception):
    logging.error(f'Exception occurred: {exception}', exc_info=True)

def notify_user(message):
    if not isinstance(message, str):
        print('Error: Message for notification is not a string.')
        return
    print(f'Notification: {message}')

def authenticate_user(username, password):
    # Check if username and password are not empty
    if not username or not password:
        print('Error: Username and password cannot be empty.')
        return False
    # This is a placeholder. In a real-world scenario, you'd have a secure way to store and check user credentials.
    if username == 'admin' and password == 'password':
        return True
    return False

import requests

def get_data_from_api(api_url):
    try:
        response = requests.get(api_url)
        response.raise_for_status()
    except requests.exceptions.HTTPError as http_err:
        print(f'HTTP error occurred: {http_err}')
        return None
    except Exception as err:
        print(f'Other error occurred: {err}')
        return None
    else:
        return response.json()
